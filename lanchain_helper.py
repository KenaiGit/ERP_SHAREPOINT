import os
import requests
from msal import ConfidentialClientApplication
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema.document import Document
import streamlit as st
from io import BytesIO
from docx import Document as DocxDocument
import PyPDF2

# 🔐 Microsoft App Credentials
CLIENT_ID = st.secrets["CLIENT_ID"]
CLIENT_SECRET = st.secrets["CLIENT_SECRET"]
TENANT_ID = st.secrets["TENANT_ID"]
AUTHORITY = f"https://login.microsoftonline.com/{TENANT_ID}"
SCOPES = ["https://graph.microsoft.com/.default"]

# 🌐 SharePoint Info
SHAREPOINT_HOST = st.secrets["SHAREPOINT_HOST"]
SITE_NAME = st.secrets["SITE_NAME"]
DOC_LIB_PATH = st.secrets["DOC_LIB_PATH"]

# 🔎 Embeddings
EMBEDDINGS_MODEL = "sentence-transformers/all-mpnet-base-v2"
embeddings = HuggingFaceEmbeddings(model_name=EMBEDDINGS_MODEL)


def authenticate():
    app = ConfidentialClientApplication(
        client_id=CLIENT_ID,
        client_credential=CLIENT_SECRET,
        authority=AUTHORITY,
    )
    result = app.acquire_token_for_client(scopes=SCOPES)
    if "access_token" not in result:
        raise Exception(f"❌ Failed to acquire token: {result.get('error_description')}")
    print("✅ Successfully authenticated via App-Only flow.")
    return result["access_token"]


def extract_text_from_docx(content: bytes) -> str:
    doc = DocxDocument(BytesIO(content))
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_pdf(content: bytes) -> str:
    reader = PyPDF2.PdfReader(BytesIO(content))
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def fetch_txt_files_from_sharepoint():
    token = authenticate()
    headers = {"Authorization": f"Bearer {token}"}
    try:
        site_url = f"https://graph.microsoft.com/v1.0/sites/{SHAREPOINT_HOST}:/sites/{SITE_NAME}"
        site_id = requests.get(site_url, headers=headers).json()["id"]

        drives_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives"
        drive_id = next((d["id"] for d in requests.get(drives_url, headers=headers).json()["value"] if d["name"] == "Documents"), None)

        encoded_path = DOC_LIB_PATH.replace(" ", "%20")
        files_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{encoded_path}:/children"
        files = requests.get(files_url, headers=headers).json().get("value", [])

        docs = []
        for item in files:
            name = item["name"]
            if name.endswith((".txt", ".docx", ".pdf")):
                file_resp = requests.get(item["@microsoft.graph.downloadUrl"])
                file_resp.raise_for_status()
                content = file_resp.content

                if name.endswith(".txt"):
                    text = content.decode("utf-8")
                elif name.endswith(".docx"):
                    text = extract_text_from_docx(content)
                elif name.endswith(".pdf"):
                    text = extract_text_from_pdf(content)
                else:
                    continue

                docs.append(Document(page_content=text, metadata={
                    "source": name,
                    "full_content": text
                }))

        print(f"📄 Retrieved {len(docs)} documents from SharePoint")
        return docs

    except Exception as e:
        print(f"❌ Error fetching documents: {e}")
        return []


def index_documents():
    print("📥 Beginning indexing of SharePoint documents...")
    documents = fetch_txt_files_from_sharepoint()
    if not documents:
        raise Exception("❌ No .txt files found to index.")

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    source_to_full = {doc.metadata["source"]: doc.metadata["full_content"] for doc in documents}
    chunks = text_splitter.split_documents(documents)

    for chunk in chunks:
        source = chunk.metadata.get("source")
        chunk.metadata["full_content"] = source_to_full.get(source, "")

    vectorstore = FAISS.from_documents(chunks, embeddings)
    vectorstore.save_local("./vector_index")
    print("✅ Indexing complete and stored locally.")


def get_similar_answer_from_documents(query: str, score_threshold=1.0):
    print(f"🧐 Searching for: {query}")
    print(f"⚙️ Using score threshold: {score_threshold}")

    if not os.path.exists("./vector_index"):
        print("⚠️ Vector index missing. Initiating indexing...")
        index_documents()

    try:
        vectorstore = FAISS.load_local("./vector_index", embeddings, allow_dangerous_deserialization=True)
    except Exception as e:
        print(f"⚠️ Error loading vector index: {e}. Rebuilding index...")
        index_documents()
        vectorstore = FAISS.load_local("./vector_index", embeddings, allow_dangerous_deserialization=True)

    docs_with_scores = vectorstore.similarity_search_with_score(query, k=5)

    if not docs_with_scores:
        return "❓ Apologies, I couldn't find relevant information.", None

    for doc, score in docs_with_scores:
        print(f"📄 {doc.metadata.get('source', 'Unknown')} — Score: {score:.4f}")
        if score < score_threshold:
            full_content = doc.metadata.get("full_content", doc.page_content)
            return f"✅ **Answer:** {doc.page_content}", full_content


    return "❌ Sorry, no results were relevant based on the current threshold.", None
